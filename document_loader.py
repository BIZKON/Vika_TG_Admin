"""
Загрузчик документов для базы знаний.

Поддерживает:
- PDF файлы
- Word документы (.docx)
- Текстовые файлы (.txt, .md)
- YAML файлы (knowledge_base.yml)

Документы разбиваются на чанки для оптимального поиска.
"""

import logging
import re
from pathlib import Path
from dataclasses import dataclass

import yaml

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """Чанк документа для индексации."""
    title: str
    content: str
    category: str
    source: str  # путь к файлу или "yaml", "telegram"
    chunk_index: int = 0


class DocumentLoader:
    """
    Загрузчик документов разных форматов.

    Разбивает документы на чанки оптимального размера для embedding.
    """

    # Оптимальный размер чанка в символах (примерно 500-1000 токенов)
    CHUNK_SIZE = 2000
    CHUNK_OVERLAP = 200

    def __init__(self, documents_dir: str = "data/documents"):
        self.documents_dir = Path(documents_dir)
        self.documents_dir.mkdir(parents=True, exist_ok=True)

    # ──────────────────────────────────────────
    # Загрузка разных форматов
    # ──────────────────────────────────────────

    def load_pdf(self, file_path: str | Path) -> list[DocumentChunk]:
        """Загрузить PDF файл."""
        try:
            from pypdf import PdfReader
        except ImportError:
            logger.error("pypdf не установлен. Выполните: pip install pypdf")
            return []

        file_path = Path(file_path)
        if not file_path.exists():
            logger.error(f"Файл не найден: {file_path}")
            return []

        try:
            reader = PdfReader(file_path)
            full_text = ""

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    full_text += text + "\n\n"

            chunks = self._split_into_chunks(
                text=full_text,
                title=file_path.stem,
                category="document",
                source=str(file_path),
            )

            logger.info(f"PDF загружен: {file_path.name} ({len(chunks)} чанков)")
            return chunks

        except Exception as e:
            logger.error(f"Ошибка при загрузке PDF {file_path}: {e}")
            return []

    def load_docx(self, file_path: str | Path) -> list[DocumentChunk]:
        """Загрузить Word документ (.docx)."""
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx не установлен. Выполните: pip install python-docx")
            return []

        file_path = Path(file_path)
        if not file_path.exists():
            logger.error(f"Файл не найден: {file_path}")
            return []

        try:
            doc = Document(file_path)
            full_text = ""

            for para in doc.paragraphs:
                if para.text.strip():
                    full_text += para.text + "\n\n"

            # Также извлекаем таблицы
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        full_text += row_text + "\n"
                full_text += "\n"

            chunks = self._split_into_chunks(
                text=full_text,
                title=file_path.stem,
                category="document",
                source=str(file_path),
            )

            logger.info(f"DOCX загружен: {file_path.name} ({len(chunks)} чанков)")
            return chunks

        except Exception as e:
            logger.error(f"Ошибка при загрузке DOCX {file_path}: {e}")
            return []

    def load_text(self, file_path: str | Path) -> list[DocumentChunk]:
        """Загрузить текстовый файл (.txt, .md)."""
        file_path = Path(file_path)
        if not file_path.exists():
            logger.error(f"Файл не найден: {file_path}")
            return []

        try:
            with open(file_path, "r", encoding="utf-8") as f:
                full_text = f.read()

            chunks = self._split_into_chunks(
                text=full_text,
                title=file_path.stem,
                category="document",
                source=str(file_path),
            )

            logger.info(f"Текст загружен: {file_path.name} ({len(chunks)} чанков)")
            return chunks

        except Exception as e:
            logger.error(f"Ошибка при загрузке текста {file_path}: {e}")
            return []

    def load_yaml_kb(self, file_path: str | Path = "data/knowledge_base.yml") -> list[DocumentChunk]:
        """Загрузить базу знаний из YAML."""
        file_path = Path(file_path)
        if not file_path.exists():
            logger.warning(f"YAML файл не найден: {file_path}")
            return []

        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = yaml.safe_load(f)

            if not data or "articles" not in data:
                return []

            chunks = []
            for article in data["articles"]:
                chunks.append(DocumentChunk(
                    title=article.get("title", ""),
                    content=article.get("content", ""),
                    category=article.get("category", "faq"),
                    source="yaml",
                    chunk_index=0,
                ))

            logger.info(f"YAML загружен: {len(chunks)} статей")
            return chunks

        except Exception as e:
            logger.error(f"Ошибка при загрузке YAML {file_path}: {e}")
            return []

    def load_file(self, file_path: str | Path) -> list[DocumentChunk]:
        """Автоматически определить тип файла и загрузить."""
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return self.load_pdf(file_path)
        elif suffix == ".docx":
            return self.load_docx(file_path)
        elif suffix in (".txt", ".md", ".rst"):
            return self.load_text(file_path)
        elif suffix in (".yml", ".yaml"):
            return self.load_yaml_kb(file_path)
        else:
            logger.warning(f"Неподдерживаемый формат файла: {suffix}")
            return []

    def load_directory(
        self,
        dir_path: str | Path | None = None,
        recursive: bool = True,
    ) -> list[DocumentChunk]:
        """Загрузить все документы из директории."""
        dir_path = Path(dir_path) if dir_path else self.documents_dir

        if not dir_path.exists():
            logger.error(f"Директория не найдена: {dir_path}")
            return []

        all_chunks = []
        patterns = ["*.pdf", "*.docx", "*.txt", "*.md"]

        for pattern in patterns:
            if recursive:
                files = dir_path.rglob(pattern)
            else:
                files = dir_path.glob(pattern)

            for file_path in files:
                chunks = self.load_file(file_path)
                all_chunks.extend(chunks)

        logger.info(f"Загружено из директории {dir_path}: {len(all_chunks)} чанков")
        return all_chunks

    # ──────────────────────────────────────────
    # Chunking
    # ──────────────────────────────────────────

    def _split_into_chunks(
        self,
        text: str,
        title: str,
        category: str,
        source: str,
    ) -> list[DocumentChunk]:
        """
        Разбить текст на чанки с перекрытием.

        Стратегия:
        1. Сначала пытаемся разбить по заголовкам/секциям
        2. Если секции слишком большие - разбиваем по абзацам
        3. Если абзацы слишком большие - разбиваем по предложениям
        """
        # Очищаем текст
        text = self._clean_text(text)

        if len(text) <= self.CHUNK_SIZE:
            # Текст достаточно короткий
            return [DocumentChunk(
                title=title,
                content=text,
                category=category,
                source=source,
                chunk_index=0,
            )]

        # Пытаемся разбить по секциям (заголовки)
        sections = self._split_by_sections(text)

        chunks = []
        chunk_index = 0

        for section_title, section_text in sections:
            # Если секция слишком большая - разбиваем дальше
            if len(section_text) > self.CHUNK_SIZE:
                sub_chunks = self._split_by_size(section_text)
                for sub_text in sub_chunks:
                    chunk_title = f"{title} - {section_title}" if section_title else title
                    chunks.append(DocumentChunk(
                        title=chunk_title,
                        content=sub_text,
                        category=category,
                        source=source,
                        chunk_index=chunk_index,
                    ))
                    chunk_index += 1
            else:
                chunk_title = f"{title} - {section_title}" if section_title else title
                chunks.append(DocumentChunk(
                    title=chunk_title,
                    content=section_text,
                    category=category,
                    source=source,
                    chunk_index=chunk_index,
                ))
                chunk_index += 1

        return chunks

    def _split_by_sections(self, text: str) -> list[tuple[str, str]]:
        """Разбить текст по заголовкам/секциям."""
        # Паттерны заголовков
        header_pattern = r'^(?:#{1,3}\s+|(?:\d+\.)+\s+|[А-ЯA-Z][А-Яа-яA-Za-z\s]{5,50}:?\s*$)'

        lines = text.split('\n')
        sections = []
        current_title = ""
        current_content = []

        for line in lines:
            if re.match(header_pattern, line.strip(), re.MULTILINE):
                # Сохраняем предыдущую секцию
                if current_content:
                    content = '\n'.join(current_content).strip()
                    if content:
                        sections.append((current_title, content))

                current_title = line.strip().lstrip('#').strip()
                current_content = []
            else:
                current_content.append(line)

        # Сохраняем последнюю секцию
        if current_content:
            content = '\n'.join(current_content).strip()
            if content:
                sections.append((current_title, content))

        # Если не удалось разбить на секции - возвращаем весь текст
        if len(sections) <= 1:
            return [("", text)]

        return sections

    def _split_by_size(self, text: str) -> list[str]:
        """Разбить текст на чанки фиксированного размера с перекрытием."""
        # Сначала разбиваем по абзацам
        paragraphs = text.split('\n\n')
        paragraphs = [p.strip() for p in paragraphs if p.strip()]

        chunks = []
        current_chunk = ""

        for para in paragraphs:
            if len(current_chunk) + len(para) <= self.CHUNK_SIZE:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())

                # Если абзац сам по себе больше chunk_size - разбиваем по предложениям
                if len(para) > self.CHUNK_SIZE:
                    sentences = self._split_into_sentences(para)
                    for sent in sentences:
                        if len(current_chunk) + len(sent) <= self.CHUNK_SIZE:
                            current_chunk += sent + " "
                        else:
                            if current_chunk:
                                chunks.append(current_chunk.strip())
                            current_chunk = sent + " "
                else:
                    current_chunk = para + "\n\n"

        if current_chunk:
            chunks.append(current_chunk.strip())

        # Добавляем перекрытие
        if self.CHUNK_OVERLAP > 0 and len(chunks) > 1:
            overlapped_chunks = [chunks[0]]
            for i in range(1, len(chunks)):
                # Добавляем конец предыдущего чанка в начало текущего
                prev_end = chunks[i - 1][-self.CHUNK_OVERLAP:]
                overlapped_chunks.append(prev_end + " ... " + chunks[i])
            chunks = overlapped_chunks

        return chunks

    def _split_into_sentences(self, text: str) -> list[str]:
        """Разбить текст на предложения."""
        # Простой паттерн для разбиения по предложениям
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]

    def _clean_text(self, text: str) -> str:
        """Очистить текст от лишних символов."""
        # Убираем множественные пробелы
        text = re.sub(r' +', ' ', text)
        # Убираем множественные переносы строк
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Убираем пробелы в начале/конце строк
        text = '\n'.join(line.strip() for line in text.split('\n'))
        return text.strip()


class TelegramHistoryLoader:
    """
    Загрузчик истории из Telegram для обучения.

    Извлекает пары вопрос-ответ из реальной переписки.
    """

    def __init__(self):
        pass

    def extract_qa_pairs(
        self,
        messages: list[dict],
        moderator_id: int,
    ) -> list[DocumentChunk]:
        """
        Извлечь пары вопрос-ответ из списка сообщений.

        Args:
            messages: Список сообщений с ключами: sender_id, text, reply_to_msg_id
            moderator_id: ID модератора (Виктории)

        Returns:
            Список DocumentChunk для индексации
        """
        # Строим индекс сообщений по ID
        msg_index = {m.get("id"): m for m in messages if m.get("id")}

        chunks = []
        seen_pairs = set()

        for msg in messages:
            # Ищем ответы Виктории
            if msg.get("sender_id") != moderator_id:
                continue

            reply_to = msg.get("reply_to_msg_id")
            if not reply_to:
                continue

            # Находим исходное сообщение
            original = msg_index.get(reply_to)
            if not original:
                continue

            question = original.get("text", "").strip()
            answer = msg.get("text", "").strip()

            if not question or not answer:
                continue

            # Фильтруем короткие/бессмысленные
            if len(question) < 10 or len(answer) < 10:
                continue

            # Дедупликация
            pair_key = f"{question[:50]}:{answer[:50]}"
            if pair_key in seen_pairs:
                continue
            seen_pairs.add(pair_key)

            chunks.append(DocumentChunk(
                title=question[:100],
                content=f"Вопрос: {question}\n\nОтвет: {answer}",
                category="learned",
                source="telegram",
                chunk_index=0,
            ))

        logger.info(f"Извлечено {len(chunks)} пар вопрос-ответ из Telegram")
        return chunks


# Глобальный loader
document_loader = DocumentLoader()
